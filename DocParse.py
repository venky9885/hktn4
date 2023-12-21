#signature
import cv2
from PIL import Image
import numpy as np
import fitz 
from docx import Document
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import words
from flask import Flask, request, jsonify
from io import BytesIO
import docx2txt
import io
import sys

# from signature_detect.cropper import Cropper
# from signature_detect.extractor import Extractor
# from signature_detect.loader import Loader
# from signature_detect.judger import Judger


app = Flask(__name__)



def has_signature(image,yesPdf):
    # Convert the Pillow image to a NumPy array
    if(yesPdf):
        img_array = np.frombuffer(image.samples, dtype=np.uint8).reshape((image.h, image.w, 3))
    else:
        img_array = np.array(image)

    if len(img_array.shape) == 3 and img_array.shape[2] == 3:
        # Convert to grayscale only if the image is not already grayscale
        gray_image = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
    else:
        gray_image = img_array

    blurred_image = cv2.GaussianBlur(gray_image, (5, 5), 0)

    edges = cv2.Canny(blurred_image, 50, 150)
    signature_contours, _ = cv2.findContours(edges.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    min_signature_contour_count = 4

    has_signature = len(signature_contours) >= min_signature_contour_count

    harris_corners = cv2.cornerHarris(blurred_image, 2, 3, 0.04)

    threshold = 0.1 * harris_corners.max()
    ridge_points = np.where(harris_corners > threshold)

    min_ridge_point_count = 100

    has_thumbprint = len(ridge_points[0]) >= min_ridge_point_count
    print(len(signature_contours),has_signature ,len(ridge_points[0]),has_thumbprint,"hii")
    if has_signature and has_thumbprint:
        return "ST"
    elif has_signature:
        return "S"
    elif has_thumbprint:
        return "T"
    else:
        return "None"

def has_signature_in_pdf(pdf_path, Docum):
    doc = Docum
    result_list = []

    for page_num in range(doc.page_count):
        page = doc[page_num]
        pix = page.get_pixmap()
        
        result_list.append(has_signature(pix,True))

    return result_list

def has_signature_in_docx(docx_path, Docum):
    doc = Docum
    result_list = []

    for rel in doc.part.rels:
        if "image" in str(doc.part.rels[rel].target_ref):
            image_data = doc.part.rels[rel].target_part.blob
            result_list.append(has_signature(Image.open(io.BytesIO(image_data)),False))

    return result_list

def validateSign(document_path, Docum):
    if document_path.lower().endswith('.pdf'):
        return has_signature_in_pdf(document_path, Docum)
    elif document_path.lower().endswith('.docx'):
        return has_signature_in_docx(document_path, Docum)
    else:
        return []

#!3 check pg count and  min word count

nltk.download('punkt')
nltk.download('words')
def is_valid_document(document_path,Docum):
    min_word_count = 5
    print(1230,document_path,document_path.lower().endswith('.pdf'))
    if document_path.lower().endswith('.pdf'):
        doc = Docum
        print(123,doc)
        # fitz.open(document_path)
        if doc.page_count < 0:
            return False
        print(1234,doc.page_count)
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text()

            tokens = [word.lower() for word in word_tokenize(text) if word.isalpha()]
            print(1235,tokens)
            if len(tokens) >= min_word_count:
                return True
        return False
    elif document_path.lower().endswith('.docx'):

        doc = Docum
        if len(doc.paragraphs) < 1:
            return False  

        for paragraph in doc.paragraphs:
            text = paragraph.text

            tokens = [word.lower() for word in word_tokenize(text) if word.isalpha()]

            if len(tokens) >= min_word_count:
                return True

        return False

    else:
        return False


#!4 atlst 50


def contains_keywords(document_path, keywords,Docum):
    if document_path.lower().endswith('.pdf'):
        doc = Docum
        # fitz.open(document_path)
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text()
            if any(keyword.lower() in text.lower() for keyword in keywords):
                return True
        return False
    elif document_path.lower().endswith('.docx'):
        doc = Docum
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if any(keyword.lower() in text.lower() for keyword in keywords):
                return True
        return False
    else:
        return False
    
#! key word check


# if __name__=="__main__":
#     document_path = "path/to/your/uploaded/document.pdf" 
#     keywords_to_check = ["Loan", "Request", "Withdrawal", "Spousal", "Consent"]
#     keywrdResult = contains_keywords(document_path, keywords_to_check)
#     validateDocumentResult = is_valid_document(document_path)
#     signResult = validateSign(document_path)

#     jsonResult = {
#         "keywordsValid":keywrdResult,
#         "validateDoc": validateDocumentResult,
#         "signatureResult": signResult
#     }

#     print(jsonResult)
@app.route("/")
def hello_world():
    return "<p>Hello, World!</p>"

@app.route('/validate_document', methods=['POST'])
def validate_document():
    # Get the uploaded file from the request
    print("21112")
    uploaded_file = request.files['document']
    keywords_to_check = ["Loan","Loans", "Request", "Withdrawal", "Spousal", "Consent"]

    # Read the content of the file
    content = uploaded_file.read()
    print(type(uploaded_file),"7895")
    # Check if the document is a PDF
    if uploaded_file.filename.lower().endswith('.pdf'):
        doc = fitz.open("pdf", content)
        print(doc,"2111",doc.page_count)
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text()
            keywords_result = contains_keywords(uploaded_file.filename, keywords_to_check,doc)

            document_valid_result = is_valid_document(uploaded_file.filename,doc)

            signature_result = validateSign(uploaded_file.filename,doc)

            # Create JSON result
            json_result = {
                "keywords_valid": keywords_result,
                "document_valid": document_valid_result,
                "signature_result": signature_result
            }

            return jsonify(json_result)
    elif uploaded_file.filename.lower().endswith('.docx'):
        # dx = docx2txt.process(uploaded_file)
        # doc = fitz.open("pdf", content)
        doc = Document(io.BytesIO(content))
        # doc =  fitz.open(stream=content, filetype="docx")
        print(doc,"8555")
        for page_num in range(len(doc.paragraphs)):
            # page = doc[page_num]
            text = doc.paragraphs[page_num]
            keywords_result = contains_keywords(uploaded_file.filename, keywords_to_check,doc)
            
            document_valid_result = is_valid_document(uploaded_file.filename,doc)
            print(document_valid_result ,"8785")
            signature_result = validateSign(uploaded_file.filename,doc)

            # Create JSON result
            json_result = {
                "keywords_valid": keywords_result,
                "document_valid": document_valid_result,
                "signature_result": signature_result
            }

            return jsonify(json_result)
        else:
            return jsonify({
                "keywords_valid": False,
                "document_valid": False,
                "signature_result": False
            })

if __name__ == '__main__':
    app.run(debug=True)
